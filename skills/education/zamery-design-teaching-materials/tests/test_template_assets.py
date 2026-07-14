import csv
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from scripts.build_docx_templates import build_docx_templates

ROOT = Path(__file__).resolve().parents[1]
TEMPLATES = ROOT / "assets" / "templates"


def grid_ratios(document: Document, table_index: int = 0) -> list[float]:
    grid = document.tables[table_index]._tbl.tblGrid
    widths = [int(column.get(qn("w:w"))) for column in grid.gridCol_lst]
    total = sum(widths)
    return [round(width / total, 2) for width in widths]


def header_text(document: Document) -> str:
    values: list[str] = []
    for section in document.sections:
        values.extend(paragraph.text for paragraph in section.header.paragraphs)
        for table in section.header.tables:
            values.extend(
                paragraph.text
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            )
    return " ".join(values)


class TemplateAssetTests(unittest.TestCase):
    def test_docx_builder_creates_seven_valid_templates(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            paths = build_docx_templates(Path(directory))
            self.assertEqual({path.name for path in paths}, {
                "teacher-guide-template.docx",
                "student-worksheet-template.docx",
                "answer-key-template.docx",
                "student-workbook-template.docx",
                "exam-paper-template.docx",
                "answer-sheet-template.docx",
                "administration-guide-template.docx",
            })
            for path in paths:
                self.assertTrue(path.is_file())
                with zipfile.ZipFile(path) as archive:
                    self.assertIsNone(archive.testzip())

    def test_committed_docx_templates_use_brand_and_asymmetric_geometry(self) -> None:
        teacher = Document(TEMPLATES / "teacher-guide-template.docx")
        student = Document(TEMPLATES / "student-worksheet-template.docx")
        answers = Document(TEMPLATES / "answer-key-template.docx")
        for document in (teacher, student, answers):
            self.assertIn("zamery", header_text(document))
            self.assertEqual(document.styles["Normal"].font.name, "Arial")
        self.assertEqual(grid_ratios(teacher), [0.18, 0.12, 0.4, 0.3])
        self.assertEqual(grid_ratios(student), [0.08, 0.55, 0.37])
        self.assertEqual(grid_ratios(answers), [0.1, 0.28, 0.62])

    def test_long_form_templates_are_committed_and_registered(self) -> None:
        expected = {
            "student-workbook-template.docx",
            "exam-paper-template.docx",
            "answer-sheet-template.docx",
            "administration-guide-template.docx",
        }
        for name in expected:
            document = Document(TEMPLATES / name)
            self.assertIn("zamery", header_text(document))
        registry = (ROOT / "assets" / "artifact-template-registry.csv").read_text(encoding="utf-8")
        for name in expected:
            self.assertIn(name, registry)

    def test_presentation_template_is_valid_ooxml_and_contains_brand_colors(self) -> None:
        path = TEMPLATES / "classroom-slides-template.pptx"
        with zipfile.ZipFile(path) as archive:
            self.assertIsNone(archive.testzip())
            xml = "\n".join(
                archive.read(name).decode("utf-8", errors="ignore")
                for name in archive.namelist()
                if name.endswith(".xml")
            )
        self.assertIn("17324D", xml)
        self.assertIn("B85435", xml)
        self.assertIn("2B746F", xml)

    def test_offline_html_and_question_bank_templates_are_present(self) -> None:
        html = (TEMPLATES / "offline-html" / "index.html").read_text(encoding="utf-8")
        css = (TEMPLATES / "offline-html" / "zamery.css").read_text(encoding="utf-8")
        self.assertIn("zamery", html)
        self.assertNotIn("http://", html)
        self.assertNotIn("https://", html)
        self.assertIn("#17324D", css)
        with (TEMPLATES / "question-bank-template.csv").open("r", encoding="utf-8", newline="") as handle:
            reader = csv.reader(handle)
            self.assertEqual(
                next(reader),
                ["item_id", "objective_id", "item_type", "prompt", "options", "response_type", "expected_response_lines", "difficulty", "source_id"],
            )


if __name__ == "__main__":
    unittest.main()
