import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from scripts.long_form_documents import build_exam_pack, build_workbook


def items(count: int) -> list[dict[str, object]]:
    return [
        {
            "item_id": f"item-{index:04d}",
            "version": 1,
            "section_id": "Grammar" if index <= count // 2 else "Reading",
            "interaction": "single_choice" if index % 3 else "short_answer",
            "stem": f"Question {index}: choose or write the best answer.",
            "options": [
                {"option_id": "A", "text": "First option"},
                {"option_id": "B", "text": "Second option"},
                {"option_id": "C", "text": "Third option"},
            ],
            "answer": {"correct_option_ids": ["B"]} if index % 3 else {"accepted_answers": ["sample"]},
            "rationale": "Teacher scoring evidence.",
        }
        for index in range(1, count + 1)
    ]


class LongFormDocumentTests(unittest.TestCase):
    def test_builds_100_item_workbook_with_continuous_numbering(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "workbook.docx"
            report = build_workbook(items(100), path, title="Grade 7 Practice Workbook")
            self.assertEqual(report["item_count"], 100)
            self.assertEqual(report["numbering"], {"first": 1, "last": 100, "continuous": True})
            document = Document(path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("1.", text)
            self.assertIn("100.", text)
            with zipfile.ZipFile(path) as archive:
                self.assertIsNone(archive.testzip())

    def test_workbook_groups_interleaved_items_into_stable_sections(self) -> None:
        mixed = items(12)
        for index, item in enumerate(mixed):
            item["section_id"] = ("Vocabulary", "Writing", "Speaking")[index % 3]
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "workbook.docx"
            report = build_workbook(mixed, path, title="Mixed Skills Workbook")
            document = Document(path)
            headings = [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.style.name == "Heading 1"
            ]
            self.assertEqual(headings[:3], ["Vocabulary", "Writing", "Speaking"])
            self.assertEqual(headings.count("Vocabulary"), 1)
            self.assertEqual(headings.count("Writing"), 1)
            self.assertEqual(headings.count("Speaking"), 1)
            self.assertEqual(report["section_order"], ["Vocabulary", "Writing", "Speaking"])
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("1.", text)
            self.assertIn("12.", text)

    def test_constructed_response_space_is_one_non_splitting_block(self) -> None:
        constructed = items(1)
        constructed[0]["interaction"] = "short_answer"
        constructed[0]["answer"] = {"accepted_answers": ["sample"]}
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "workbook.docx"
            build_workbook(constructed, path, title="Constructed Response Workbook")
            document = Document(path)
            prompt_index = next(
                index
                for index, paragraph in enumerate(document.paragraphs)
                if paragraph.text.startswith("1.")
            )
            response = document.paragraphs[prompt_index + 1]
            self.assertEqual(response.text.count("_") , 80 * 3)
            self.assertEqual(response.text.count("\n"), 2)
            self.assertTrue(response.paragraph_format.keep_together)

    def test_exam_pack_keeps_student_teacher_and_admin_surfaces_separate(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            paths, report = build_exam_pack(items(80), Path(directory), title="Grade 8 Midterm")
            self.assertEqual(set(paths), {"exam_paper", "answer_sheet", "answer_key", "administration_guide"})
            self.assertEqual(report["item_count"], 80)
            exam_text = "\n".join(paragraph.text for paragraph in Document(paths["exam_paper"]).paragraphs).casefold()
            key_text = "\n".join(paragraph.text for paragraph in Document(paths["answer_key"]).paragraphs).casefold()
            self.assertNotIn("teacher scoring evidence", exam_text)
            self.assertIn("teacher scoring evidence", key_text)
            for path in paths.values():
                with zipfile.ZipFile(path) as archive:
                    self.assertIsNone(archive.testzip())


if __name__ == "__main__":
    unittest.main()
