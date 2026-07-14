from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = "17324D"
TERRACOTTA = "B85435"
TEAL = "2B746F"
SAND = "F4EEE4"
MIST = "DCEAF2"
CHARCOAL = "202B33"
WHITE = "FFFFFF"
USABLE_DXA = 10224  # 7.1 in at 0.7 in side margins on US Letter.


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != USABLE_DXA:
        raise ValueError("table widths must sum to usable width")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def _style_table(table, header_fill: str = MIST) -> None:
    table.style = "Table Grid"
    _set_repeat_header(table.rows[0])
    for index, cell in enumerate(table.rows[0].cells):
        _set_cell_shading(cell, header_fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(INK)
                run.font.name = "Arial"
                run.font.size = Pt(9.5)
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)


def _paragraph_shading(paragraph, fill: str, border_color: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    if border_color:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        borders.append(left)
        p_pr.append(borders)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("zamery  •  rooted in strength     ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(INK)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _configure_document(document: Document, audience: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.right_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.3)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(CHARCOAL)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in (
        ("Title", 24, INK, 0, 10),
        ("Heading 1", 15, INK, 14, 6),
        ("Heading 2", 12, TEAL, 10, 4),
    ):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header.paragraphs[0].paragraph_format.space_after = Pt(0)
    table = header.add_table(rows=1, cols=2, width=Inches(7.1))
    _set_table_geometry(table, [6500, 3724])
    for cell in table.rows[0].cells:
        _set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    brand_paragraph = table.rows[0].cells[0].paragraphs[0]
    brand_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    wordmark = brand_paragraph.add_run("zamery")
    wordmark.bold = True
    wordmark.font.name = "Arial"
    wordmark.font.size = Pt(13)
    wordmark.font.color.rgb = RGBColor.from_string(INK)
    tagline = brand_paragraph.add_run("   rooted in strength")
    tagline.italic = True
    tagline.font.name = "Arial"
    tagline.font.size = Pt(8)
    tagline.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    audience_paragraph = table.rows[0].cells[1].paragraphs[0]
    audience_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    audience_run = audience_paragraph.add_run(audience.upper())
    audience_run.bold = True
    audience_run.font.name = "Arial"
    audience_run.font.size = Pt(8)
    audience_run.font.color.rgb = RGBColor.from_string(TEAL)

    _add_page_number(section.footer.paragraphs[0])


def _add_opening(document: Document, eyebrow: str, title: str, subtitle: str) -> None:
    eyebrow_paragraph = document.add_paragraph()
    eyebrow_run = eyebrow_paragraph.add_run(eyebrow.upper())
    eyebrow_run.bold = True
    eyebrow_run.font.name = "Arial"
    eyebrow_run.font.size = Pt(9)
    eyebrow_run.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    eyebrow_paragraph.paragraph_format.space_after = Pt(4)
    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.add_run(title)
    subtitle_paragraph = document.add_paragraph()
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor.from_string(TEAL)
    subtitle_paragraph.paragraph_format.space_after = Pt(12)


def _add_callout(document: Document, label: str, text: str, fill: str = SAND, accent: str = TERRACOTTA) -> None:
    paragraph = document.add_paragraph()
    _paragraph_shading(paragraph, fill, accent)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    label_run = paragraph.add_run(f"{label.upper()}  ")
    label_run.bold = True
    label_run.font.name = "Arial"
    label_run.font.size = Pt(9.5)
    label_run.font.color.rgb = RGBColor.from_string(accent)
    text_run = paragraph.add_run(text)
    text_run.font.name = "Arial"
    text_run.font.size = Pt(10.5)
    text_run.font.color.rgb = RGBColor.from_string(CHARCOAL)


def _add_bullets(document: Document, values: list[str]) -> None:
    for value in values:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(value)
        paragraph.paragraph_format.space_after = Pt(3)


def build_teacher_template(path: Path) -> None:
    document = Document()
    _configure_document(document, "Teacher guide")
    _add_opening(document, "Lesson command center", "{{ lesson_title }}", "{{ grade }}  •  {{ duration }}  •  {{ instruction_language }}")
    _add_callout(document, "Success evidence", "{{ what students will say, write, or do by the end }}", MIST, TEAL)
    document.add_heading("Objectives", level=1)
    _add_bullets(document, ["{{ objective_1 }}", "{{ objective_2 }}", "{{ objective_3 }}"])
    document.add_heading("Ready before students arrive", level=1)
    _add_bullets(document, ["Materials: {{ materials }}", "Preparation: {{ preparation }}"])
    _add_callout(document, "Watch for", "{{ highest-value misconception and recovery cue }}")

    document.add_heading("Teach — minute by minute", level=1)
    table = document.add_table(rows=1, cols=4)
    for cell, text in zip(table.rows[0].cells, ("Stage", "Min", "Teacher move", "Student evidence")):
        cell.text = text
    for stage in ("Launch", "Notice", "Test the clue", "Build the rule", "Transfer"):
        cells = table.add_row().cells
        cells[0].text = stage
        cells[1].text = "{{ min }}"
        cells[2].text = "{{ concise teacher move }}"
        cells[3].text = "{{ observable response }}"
    _set_table_geometry(table, [1840, 1227, 4090, 3067])
    _style_table(table)
    document.add_heading("Board plan", level=1)
    _add_callout(document, "Board", "{{ minimal board layout, examples, and final rule }}", MIST, TEAL)

    document.add_page_break()
    document.add_heading("Respond in the moment", level=1)
    document.add_heading("Checks for understanding", level=2)
    _add_bullets(document, ["{{ hinge question }}", "{{ fast evidence check }}"])
    document.add_heading("Support", level=2)
    _add_callout(document, "If students need support", "{{ scaffold, model, or language frame }}", MIST, TEAL)
    document.add_heading("Challenge", level=2)
    _add_callout(document, "If students are ready", "{{ deeper contrast or transfer prompt }}", SAND, TERRACOTTA)
    document.add_heading("Recovery path", level=2)
    _add_bullets(document, ["{{ likely wrong answer → probing question → repair move }}"])

    document.add_heading("Answers — teacher only", level=1)
    answer_table = document.add_table(rows=1, cols=3)
    for cell, text in zip(answer_table.rows[0].cells, ("Item", "Answer", "Clue / rationale")):
        cell.text = text
    for index in range(1, 6):
        cells = answer_table.add_row().cells
        cells[0].text = f"{{{{ item_{index} }}}}"
        cells[1].text = "{{ answer }}"
        cells[2].text = "{{ clue, accepted alternative, or rationale }}"
    _set_table_geometry(answer_table, [1022, 2863, 6339])
    _style_table(answer_table, SAND)
    document.save(path)


def build_student_template(path: Path) -> None:
    document = Document()
    _configure_document(document, "Student worksheet")
    _add_opening(document, "Clue detective", "{{ worksheet_title }}", "Name ____________________   Class __________   Date __________")
    _add_callout(document, "Your mission", "{{ direct student-facing instruction }}", MIST, TEAL)
    document.add_heading("1 · Notice", level=1)
    document.add_paragraph("{{ short bad example or contrast pair }}")
    document.add_heading("2 · Test the clue", level=1)
    table = document.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ("#", "Prompt", "Your thinking")):
        cell.text = text
    for index in range(1, 5):
        cells = table.add_row().cells
        cells[0].text = str(index)
        cells[1].text = "{{ concise prompt }}"
        cells[2].text = "{{ response line }}"
    _set_table_geometry(table, [818, 5623, 3783])
    _style_table(table)

    document.add_page_break()
    document.add_heading("3 · Explain", level=1)
    document.add_paragraph("{{ explanation prompt }}")
    for _ in range(4):
        document.add_paragraph("________________________________________________________________________________")
    document.add_heading("4 · Transfer", level=1)
    document.add_paragraph("{{ two-sentence transfer prompt }}")
    for _ in range(5):
        document.add_paragraph("________________________________________________________________________________")
    document.add_heading("5 · Exit", level=1)
    exit_table = document.add_table(rows=1, cols=2)
    exit_table.rows[0].cells[0].text = "Quick check"
    exit_table.rows[0].cells[1].text = "Answer"
    for _ in range(3):
        cells = exit_table.add_row().cells
        cells[0].text = "{{ exit item }}"
        cells[1].text = "{{ response }}"
    _set_table_geometry(exit_table, [6134, 4090])
    _style_table(exit_table, MIST)
    document.save(path)


def build_answer_template(path: Path) -> None:
    document = Document()
    _configure_document(document, "Teacher only")
    _add_opening(document, "Teacher only", "{{ answer_key_title }}", "Source: {{ source_artifact_id }}  •  Version {{ source_version }}")
    _add_callout(document, "Marking principle", "{{ answer authority and accepted alternatives }}", SAND, TERRACOTTA)
    table = document.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ("Item", "Answer", "Clue / rationale")):
        cell.text = text
    for index in range(1, 11):
        cells = table.add_row().cells
        cells[0].text = str(index)
        cells[1].text = "{{ answer }}"
        cells[2].text = "{{ evidence, accepted alternative, or misconception note }}"
    _set_table_geometry(table, [1022, 2863, 6339])
    _style_table(table, SAND)
    document.save(path)


def build_workbook_template(path: Path) -> None:
    document = Document()
    _configure_document(document, "Student workbook")
    _add_opening(document, "Practice journey", "{{ workbook_title }}", "Name ____________________   Class __________   Term __________")
    _add_callout(document, "How to use this book", "Complete one section at a time. Mark a checkpoint after review, then return to missed items.", MIST, TEAL)
    document.add_heading("Your progress map", level=1)
    table = document.add_table(rows=1, cols=4)
    for cell, text in zip(table.rows[0].cells, ("Section", "Goal", "Items", "Checkpoint")):
        cell.text = text
    for index in range(1, 5):
        cells = table.add_row().cells
        cells[0].text = f"{{{{ section_{index} }}}}"
        cells[1].text = "{{ student-friendly goal }}"
        cells[2].text = "{{ item range }}"
        cells[3].text = "□ Ready  □ Review"
    _set_table_geometry(table, [2045, 4090, 1840, 2249])
    _style_table(table)
    document.add_page_break()
    document.add_heading("Section {{ n }} · {{ section_title }}", level=1)
    _add_callout(document, "Target", "{{ what success looks like in this section }}", MIST, TEAL)
    for index in range(1, 7):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_with_next = True
        number = paragraph.add_run(f"{{{{ item_{index}_number }}}}.  ")
        number.bold = True
        number.font.color.rgb = RGBColor.from_string(TERRACOTTA)
        paragraph.add_run("{{ prompt }}")
        document.add_paragraph("{{ option or response surface }}")
    document.add_heading("Checkpoint", level=2)
    _add_callout(document, "Reflect", "One pattern I can now explain: ______________________________________________", SAND, TERRACOTTA)
    document.save(path)


def build_exam_template(path: Path) -> None:
    document = Document()
    _configure_document(document, "Student exam")
    _add_opening(document, "Assessment", "{{ exam_title }}", "{{ form_id }}  •  {{ duration }}  •  {{ total_marks }} marks")
    identity = document.add_table(rows=2, cols=4)
    values = (("Name", "________________________", "Class", "____________"), ("Candidate ID", "________________________", "Date", "____________"))
    for row, row_values in zip(identity.rows, values):
        for cell, text in zip(row.cells, row_values):
            cell.text = text
    _set_table_geometry(identity, [1431, 3681, 1431, 3681])
    _style_table(identity, SAND)
    _add_callout(document, "Instructions", "{{ permitted materials, timing, response method, and checking rule }}", MIST, TEAL)
    document.add_heading("Section 1 · {{ section_title }}", level=1)
    section_table = document.add_table(rows=1, cols=3)
    for cell, text in zip(section_table.rows[0].cells, ("Item", "Question", "Marks")):
        cell.text = text
    for index in range(1, 6):
        cells = section_table.add_row().cells
        cells[0].text = f"{{{{ item_{index}_number }}}}"
        cells[1].text = "{{ prompt and response surface }}"
        cells[2].text = "{{ mark }}"
    _set_table_geometry(section_table, [1022, 7975, 1227])
    _style_table(section_table)
    _add_callout(document, "End of section", "Check that every response is recorded on the correct line or answer sheet.", SAND, TERRACOTTA)
    document.save(path)


def build_answer_sheet_template(path: Path) -> None:
    document = Document()
    _configure_document(document, "Student answer sheet")
    _add_opening(document, "Answer sheet", "{{ exam_title }}", "{{ form_id }}  •  Write clearly and use the item number shown.")
    _add_callout(document, "Candidate", "Name ____________________   Candidate ID ____________________   Class __________", MIST, TEAL)
    document.add_heading("Selected responses", level=1)
    table = document.add_table(rows=1, cols=4)
    for cell, text in zip(table.rows[0].cells, ("Item", "Response", "Item", "Response")):
        cell.text = text
    for index in range(1, 26):
        cells = table.add_row().cells
        cells[0].text = f"{{{{ item_{index} }}}}"
        cells[1].text = "A  B  C  D"
        cells[2].text = f"{{{{ item_{index + 25} }}}}"
        cells[3].text = "A  B  C  D"
    _set_table_geometry(table, [1022, 4090, 1022, 4090])
    _style_table(table)
    document.add_page_break()
    document.add_heading("Constructed responses", level=1)
    for index in range(1, 5):
        document.add_paragraph(f"{{{{ response_item_{index} }}}}.  {{{{ prompt_short }}}}", style="Heading 2")
        for _ in range(5):
            document.add_paragraph("________________________________________________________________________________")
    document.save(path)


def build_administration_template(path: Path) -> None:
    document = Document()
    _configure_document(document, "Administration guide")
    _add_opening(document, "Teacher only", "{{ exam_title }} · Administration guide", "{{ form_ids }}  •  {{ duration }}  •  {{ item_count }} items")
    _add_callout(document, "Security", "Keep student papers, answer sheets, keys, and form manifests separated before and after administration.", SAND, TERRACOTTA)
    document.add_heading("Before students enter", level=1)
    _add_bullets(document, ["Verify form labels and page counts.", "Prepare accommodations without changing the construct.", "Check permitted materials and room setup."])
    document.add_heading("Timing and script", level=1)
    table = document.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ("Time", "Read / do", "Record")):
        cell.text = text
    for label in ("Start", "Section change", "10 minutes left", "Stop"):
        cells = table.add_row().cells
        cells[0].text = "{{ time }}"
        cells[1].text = f"{label}: {{{{ standardized script }}}}"
        cells[2].text = "{{ incident / absence }}"
    _set_table_geometry(table, [1840, 5521, 2863])
    _style_table(table)
    document.add_heading("After the assessment", level=1)
    _add_bullets(document, ["Count and reconcile every paper.", "Store answer data separately from student files.", "Record incidents before scoring."])
    document.add_heading("Incident log", level=1)
    _add_callout(document, "Record", "Candidate __________  Time ______  Event ____________________  Action ____________________", MIST, TEAL)
    document.save(path)


def build_docx_templates(output_dir: Path) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    paths = [
        output_dir / "teacher-guide-template.docx",
        output_dir / "student-worksheet-template.docx",
        output_dir / "answer-key-template.docx",
        output_dir / "student-workbook-template.docx",
        output_dir / "exam-paper-template.docx",
        output_dir / "answer-sheet-template.docx",
        output_dir / "administration-guide-template.docx",
    ]
    build_teacher_template(paths[0])
    build_student_template(paths[1])
    build_answer_template(paths[2])
    build_workbook_template(paths[3])
    build_exam_template(paths[4])
    build_answer_sheet_template(paths[5])
    build_administration_template(paths[6])
    return paths


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    for path in build_docx_templates(args.output_dir):
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
