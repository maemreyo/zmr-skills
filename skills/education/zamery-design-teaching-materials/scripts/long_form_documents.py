from __future__ import annotations

import argparse
import json
import math
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from scripts.build_docx_templates import (
    MIST,
    SAND,
    TEAL,
    TERRACOTTA,
    _add_bullets,
    _add_callout,
    _add_opening,
    _configure_document,
    _set_table_geometry,
    _style_table,
)


def _question(document: Document, number: int, item: dict[str, object], *, exam: bool) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"{number}.  ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    paragraph.add_run(str(item.get("stem", "")))
    if item.get("interaction") in {"single_choice", "multiple_choice", "true_false"}:
        options = item.get("options", [])
        option_text = "    ".join(f"{option.get('option_id')}. {option.get('text')}" for option in options if isinstance(option, dict))
        option_paragraph = document.add_paragraph(option_text)
        option_paragraph.paragraph_format.left_indent = Inches(0.25)
        option_paragraph.paragraph_format.space_after = Pt(6)
    else:
        lines = 2 if exam else 3
        response = document.add_paragraph(
            "\n".join("________________________________________________________________________________" for _ in range(lines))
        )
        response.paragraph_format.keep_together = True
        response.paragraph_format.space_after = Pt(1)


def _numbering_report(items: list[dict[str, object]]) -> dict[str, object]:
    count = len(items)
    return {"first": 1 if count else None, "last": count if count else None, "continuous": True}


def _group_items_by_section(items: list[dict[str, object]]) -> tuple[list[dict[str, object]], list[str]]:
    """Keep first-seen section order and preserve item order inside each section."""
    groups: dict[str, list[dict[str, object]]] = {}
    for item in items:
        section = str(item.get("section_id", "Practice"))
        groups.setdefault(section, []).append(item)
    section_order = list(groups)
    return [item for section in section_order for item in groups[section]], section_order


def build_workbook(items: list[dict[str, object]], path: Path, *, title: str) -> dict[str, object]:
    document = Document()
    _configure_document(document, "Student workbook")
    _add_opening(document, "Practice journey", title, "Name ____________________   Class __________   Date started __________")
    _add_callout(document, "Plan", f"{len(items)} items · Review at each section checkpoint · Return to missed items.", MIST, TEAL)
    arranged_items, section_order = _group_items_by_section(items)
    current_section: object = None
    items_on_page = 0
    for number, item in enumerate(arranged_items, 1):
        section = item.get("section_id", "Practice")
        if section != current_section:
            if current_section is not None:
                document.add_page_break()
            document.add_heading(f"{section}", level=1)
            _add_callout(document, "Section target", "Read the directions, show your thinking, and check your response.", SAND, TERRACOTTA)
            current_section = section
            items_on_page = 0
        elif items_on_page >= 20:
            document.add_page_break()
            document.add_heading(f"{section} · continued", level=2)
            items_on_page = 0
        _question(document, number, item, exam=False)
        items_on_page += 1
    document.add_page_break()
    document.add_heading("Progress reflection", level=1)
    _add_callout(document, "Reflect", "Use this page after checking your work. Write specific evidence, not only a score.", MIST, TEAL)
    for prompt, lines in (
        ("One pattern I can now explain", 3),
        ("Three items I will revisit and why", 4),
        ("My next practice goal", 3),
    ):
        document.add_heading(prompt, level=2)
        for _ in range(lines):
            document.add_paragraph("________________________________________________________________________________")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return {
        "artifact_type": "student_workbook",
        "item_count": len(items),
        "numbering": _numbering_report(items),
        "section_order": section_order,
        "planned_pages_min": max(1, math.ceil(len(items) / 20)),
    }


def _build_exam(items: list[dict[str, object]], path: Path, title: str) -> None:
    document = Document()
    _configure_document(document, "Student exam")
    _add_opening(document, "Assessment", title, f"{len(items)} items  •  Name ____________________  •  Candidate ID __________")
    _add_callout(document, "Instructions", "Answer every item. Record selected responses clearly. Check item numbers before submitting.", MIST, TEAL)
    current_section: object = None
    items_on_page = 0
    for number, item in enumerate(items, 1):
        section = item.get("section_id", "Assessment")
        if section != current_section or items_on_page >= 20:
            if current_section is not None:
                document.add_page_break()
            label = str(section) if section != current_section else f"{section} · continued"
            document.add_heading(label, level=1)
            current_section = section
            items_on_page = 0
        _question(document, number, item, exam=True)
        items_on_page += 1
    document.save(path)


def _build_answer_sheet(items: list[dict[str, object]], path: Path, title: str) -> None:
    document = Document()
    _configure_document(document, "Student answer sheet")
    _add_opening(document, "Answer sheet", title, "Name ____________________   Candidate ID ____________________")
    _add_callout(document, "Record", "Write the item number and one clear response. Continue on a labeled page if directed.", MIST, TEAL)
    for start in range(0, len(items), 30):
        if start:
            document.add_page_break()
            document.add_heading("Answer sheet · continued", level=1)
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Item"
        table.rows[0].cells[1].text = "Response"
        for number in range(start + 1, min(start + 31, len(items) + 1)):
            cells = table.add_row().cells
            cells[0].text = str(number)
            cells[1].text = "A  B  C  D   /   __________________________________"
        _set_table_geometry(table, [1227, 8997])
        _style_table(table)
    document.save(path)


def _answer_text(item: dict[str, object]) -> str:
    answer = item.get("answer", {})
    if isinstance(answer, dict) and answer.get("correct_option_ids"):
        return ", ".join(str(value) for value in answer["correct_option_ids"])
    if isinstance(answer, dict) and answer.get("accepted_answers"):
        return "; ".join(str(value) for value in answer["accepted_answers"])
    return "Rubric required"


def _build_answer_key(items: list[dict[str, object]], path: Path, title: str) -> None:
    document = Document()
    _configure_document(document, "Teacher only")
    _add_opening(document, "Teacher only", f"{title} · Answer key", f"{len(items)} items · Keep separate from student files")
    if items:
        _add_callout(document, "Teacher scoring evidence", str(items[0].get("rationale", "Use the approved answer authority.")), SAND, TERRACOTTA)
    table = document.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ("Item", "Answer", "Evidence / rationale")):
        cell.text = text
    for number, item in enumerate(items, 1):
        cells = table.add_row().cells
        cells[0].text = str(number)
        cells[1].text = _answer_text(item)
        cells[2].text = str(item.get("rationale", ""))
    _set_table_geometry(table, [1022, 2863, 6339])
    _style_table(table, SAND)
    document.save(path)


def _build_admin(items: list[dict[str, object]], path: Path, title: str) -> None:
    document = Document()
    _configure_document(document, "Administration guide")
    _add_opening(document, "Teacher only", f"{title} · Administration guide", f"{len(items)} items · Verify form and page count before release")
    _add_callout(document, "Security", "Keep exam papers, answer sheets, and answer keys physically or digitally separated.", SAND, TERRACOTTA)
    document.add_heading("Before", level=1)
    _add_bullets(document, ["Verify candidate materials and accommodations.", "Count every paper.", "Read only the standardized instructions."])
    document.add_heading("During", level=1)
    _add_bullets(document, ["Record start/stop times.", "Log incidents without coaching.", "Announce the remaining time consistently."])
    document.add_heading("After", level=1)
    _add_bullets(document, ["Reconcile every paper and answer sheet.", "Record absences and incidents.", "Release teacher scoring files only to authorized staff."])
    document.save(path)


def build_exam_pack(
    items: list[dict[str, object]], output_dir: Path, *, title: str
) -> tuple[dict[str, Path], dict[str, object]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    paths = {
        "exam_paper": output_dir / "exam-paper.docx",
        "answer_sheet": output_dir / "answer-sheet.docx",
        "answer_key": output_dir / "answer-key.docx",
        "administration_guide": output_dir / "administration-guide.docx",
    }
    _build_exam(items, paths["exam_paper"], title)
    _build_answer_sheet(items, paths["answer_sheet"], title)
    _build_answer_key(items, paths["answer_key"], title)
    _build_admin(items, paths["administration_guide"], title)
    return paths, {
        "artifact_type": "exam_pack",
        "item_count": len(items),
        "numbering": _numbering_report(items),
        "surfaces": list(paths),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Build long-form Zamery workbooks or exam packs from item JSON.")
    parser.add_argument("mode", choices=("workbook", "exam"))
    parser.add_argument("items", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--title", required=True)
    args = parser.parse_args()
    records = json.loads(args.items.read_text(encoding="utf-8"))
    if args.mode == "workbook":
        report = build_workbook(records, args.output, title=args.title)
    else:
        _, report = build_exam_pack(records, args.output, title=args.title)
    print(json.dumps(report, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
